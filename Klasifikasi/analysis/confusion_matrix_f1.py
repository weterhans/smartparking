"""
Confusion Matrix + F1 Score dari hasil uji coba klasifikasi kendaraan
Data: 24 sampel uji, 5 kelas kendaraan
Ground truth: Jenis Kendaraan
Prediksi:     Hasil Deteksi
"""

import numpy as np
import matplotlib.pyplot as plt
from matplotlib.colors import LinearSegmentedColormap
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# DATA UJI COBA
# Format: (ground_truth, prediksi)
# ============================================================
data_uji = [
    # MPV
    ("MPV",      "MPV"),       # 1. Mobilio
    ("MPV",      "MPV"),       # 2. Mobilio
    ("MPV",      "MPV"),       # 3. Xpander
    ("MPV",      "MPV"),       # 4. Avanza
    ("MPV",      "City Car"),  # 5. Grandlivina  <-- misclassification

    # SUV
    ("SUV",      "MPV"),       # 6. Pajero       <-- misclassification
    ("SUV",      "SUV"),       # 7. HRV
    ("SUV",      "SUV"),       # 8. HRV
    ("SUV",      "SUV"),       # 9. Rush
    ("SUV",      "SUV"),       # 10. Rush

    # City Car
    ("City Car", "City Car"),  # 11. Jazz
    ("City Car", "City Car"),  # 12. Brio
    ("City Car", "City Car"),  # 13. Picanto
    ("City Car", "City Car"),  # 14. Picanto
    ("City Car", "City Car"),  # 15. Picanto

    # Pickup
    ("Pickup",   "Pickup"),    # 16. Colt L300
    ("Pickup",   "Pickup"),    # 17. Colt L300
    ("Pickup",   "Pickup"),    # 18. Colt L300
    ("Pickup",   "Pickup"),    # 19. Colt L300
    ("Pickup",   "Pickup"),    # 20. Colt L300

    # Sedan
    ("Sedan",    "Sedan"),     # 21. Accord
    ("Sedan",    "Sedan"),     # 22. Accord
    ("Sedan",    "Sedan"),     # 23. Accord
    ("Sedan",    "Sedan"),     # 24. Accord (dua pengujian terpisah, nilai identik)
]

# ============================================================
# URUTAN KELAS
# ============================================================
classes = ["MPV", "SUV", "City Car", "Pickup", "Sedan"]
n = len(classes)
N = len(data_uji)

# ============================================================
# BANGUN CONFUSION MATRIX
# Baris = True label, Kolom = Predicted label
# ============================================================
cm = np.zeros((n, n), dtype=int)
idx = {c: i for i, c in enumerate(classes)}

for (true, pred) in data_uji:
    cm[idx[true]][idx[pred]] += 1

print("=" * 60)
print("CONFUSION MATRIX")
print("=" * 60)
top_left_lbl = "True \\ Pred"
header = f"{top_left_lbl:<12}" + "".join(f"{c:>12}" for c in classes)
print(header)
for i, row_label in enumerate(classes):
    row_str = f"{row_label:<12}" + "".join(f"{cm[i][j]:>12}" for j in range(n))
    print(row_str)

# ============================================================
# HITUNG TP, FP, FN, TN per KELAS + METRIK
# ============================================================
metrics = {}
support = {}
for i, cls in enumerate(classes):
    TP = int(cm[i][i])
    FP = int(cm[:, i].sum() - TP)
    FN = int(cm[i, :].sum() - TP)
    TN = int(N - TP - FP - FN)
    support[cls] = int(cm[i, :].sum())

    precision = TP / (TP + FP) if (TP + FP) > 0 else 0.0
    recall    = TP / (TP + FN) if (TP + FN) > 0 else 0.0
    f1        = (2 * precision * recall) / (precision + recall) if (precision + recall) > 0 else 0.0
    accuracy  = (TP + TN) / N

    metrics[cls] = {
        "TP": TP, "FP": FP, "FN": FN, "TN": TN,
        "Precision": precision,
        "Recall":    recall,
        "F1":        f1,
        "Accuracy":  accuracy,
    }

total_support = N
macro_f1    = float(np.mean([metrics[c]["F1"] for c in classes]))
weighted_f1 = float(sum(metrics[c]["F1"] * support[c] for c in classes) / total_support)
overall_acc = float(cm.diagonal().sum() / N)

print("\n" + "=" * 60)
print("METRIK PER KELAS")
print("=" * 60)
print(f"{'Kelas':<12} {'Precision':>10} {'Recall':>8} {'F1-Score':>10} {'Support':>8}")
for cls in classes:
    m = metrics[cls]
    print(f"{cls:<12} {m['Precision']:>10.4f} {m['Recall']:>8.4f} {m['F1']:>10.4f} {support[cls]:>8}")
print("-" * 60)
print(f"{'Macro Avg':<12} {'':>10} {'':>8} {macro_f1:>10.4f} {total_support:>8}")
print(f"{'Weighted Avg':<12} {'':>10} {'':>8} {weighted_f1:>10.4f} {total_support:>8}")
print(f"\nOverall Accuracy : {overall_acc:.4f}  ({overall_acc*100:.2f}%)")


# ============================================================
# PLOT 1 — CONFUSION MATRIX (PNG)
# ============================================================
def plot_confusion_matrix(cm, classes, filename="Confusion_Matrix_F1.png"):
    fig, ax = plt.subplots(figsize=(8, 6.5))

    cmap = LinearSegmentedColormap.from_list(
        "skripsi_blue", ["#f0f4ff", "#1a4d8f"], N=256
    )
    im = ax.imshow(cm, interpolation="nearest", cmap=cmap, vmin=0, vmax=max(cm.max(), 1))
    cbar = ax.figure.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    cbar.ax.tick_params(labelsize=9)

    ax.set_xticks(np.arange(n))
    ax.set_yticks(np.arange(n))
    ax.set_xticklabels(classes, fontsize=10, fontweight="bold")
    ax.set_yticklabels(classes, fontsize=10, fontweight="bold")
    plt.setp(ax.get_xticklabels(), rotation=20, ha="right", rotation_mode="anchor")

    ax.set_ylabel("True Label", fontsize=12, labelpad=10)
    ax.set_xlabel("Predicted Label", fontsize=12, labelpad=10)
    ax.set_title(
        "Confusion Matrix Klasifikasi Kendaraan\n(Hasil Uji Coba — 24 Data)",
        fontsize=13, fontweight="bold", pad=16
    )

    thresh = cm.max() / 2.0
    for i in range(n):
        for j in range(n):
            color = "white" if cm[i, j] > thresh else "#1a3c6e"
            weight = "bold" if i == j else "normal"
            ax.text(j, i, str(cm[i, j]),
                    ha="center", va="center",
                    color=color, fontsize=13, fontweight=weight)

    ax.set_xticks(np.arange(n) - 0.5, minor=True)
    ax.set_yticks(np.arange(n) - 0.5, minor=True)
    ax.grid(which="minor", color="white", linewidth=1.5)
    ax.tick_params(which="minor", bottom=False, left=False)

    fig.tight_layout()
    plt.savefig(filename, dpi=200, bbox_inches="tight")
    plt.close()
    print(f"[OK] Confusion matrix disimpan: {filename}")
    return filename


# ============================================================
# PLOT 2 — TABEL F1 SCORE (PNG)
# ============================================================
def plot_f1_table(metrics, classes, support, macro_f1, weighted_f1, overall_acc,
                  filename="Tabel_F1_Score.png"):
    fig, ax = plt.subplots(figsize=(9, 4.2))
    ax.axis("off")

    col_labels = ["Kelas", "Precision", "Recall", "F1-Score", "Support"]
    rows = []
    for cls in classes:
        m = metrics[cls]
        rows.append([
            cls,
            f"{m['Precision']:.4f}",
            f"{m['Recall']:.4f}",
            f"{m['F1']:.4f}",
            str(support[cls])
        ])
    rows.append(["", "", "", "", ""])
    rows.append(["Macro Avg",    "-", "-", f"{macro_f1:.4f}",    str(total_support)])
    rows.append(["Weighted Avg", "-", "-", f"{weighted_f1:.4f}", str(total_support)])

    table = ax.table(
        cellText=rows,
        colLabels=col_labels,
        cellLoc="center",
        loc="center"
    )
    table.auto_set_font_size(False)
    table.set_fontsize(10)
    table.scale(1, 1.60)

    for j in range(5):
        table[0, j].set_facecolor("#1a4d8f")
        table[0, j].set_text_props(color="white", fontweight="bold")

    for i, row in enumerate(rows):
        for j in range(5):
            if i < len(classes):
                bg = "#dde9f7" if i % 2 == 0 else "#f5f8ff"
            elif row[0] == "":
                bg = "#ffffff"
            elif row[0] == "Macro Avg":
                bg = "#c9dcf0"
            else:
                bg = "#b0cceb"
            table[i + 1, j].set_facecolor(bg)

    for i in range(len(classes)):
        table[i + 1, 3].set_text_props(fontweight="bold", color="#1a3c6e")

    for i_off in [len(classes) + 1, len(classes) + 2]:
        for j in range(5):
            table[i_off + 1, j].set_text_props(fontweight="bold")

    ax.set_title(
        f"Tabel F1-Score Klasifikasi Kendaraan\n"
        f"Overall Accuracy: {overall_acc*100:.2f}%  |  "
        f"Macro F1: {macro_f1:.4f}  |  Weighted F1: {weighted_f1:.4f}",
        fontsize=11, fontweight="bold", pad=14, color="#1a3c6e"
    )

    fig.tight_layout()
    plt.savefig(filename, dpi=200, bbox_inches="tight")
    plt.close()
    print(f"[OK] Tabel F1 disimpan: {filename}")
    return filename


# ============================================================
# HELPER WORD
# ============================================================
def set_cell_bg(cell, color_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, center=True, font_size=11,
                  color_rgb=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "Times New Roman"
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# ============================================================
# EXPORT KE WORD (.docx)
# ============================================================
def export_word(cm, classes, metrics, support, macro_f1, weighted_f1,
                overall_acc, filename="Confusion_Matrix_F1.docx"):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # ── JUDUL ──────────────────────────────────────────────
    judul = doc.add_paragraph()
    judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = judul.add_run("Confusion Matrix Klasifikasi Kendaraan")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    # ── SUBTITLE confusion matrix ───────────────────────────
    sub1 = doc.add_paragraph()
    sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = sub1.add_run("Tabel 1. Confusion Matrix (24 Data Uji)")
    r1.bold = True
    r1.font.name = "Times New Roman"

    nk = len(classes)
    tbl = doc.add_table(rows=nk + 2, cols=nk + 2)
    tbl.style = "Table Grid"

    # Pojok kiri atas (kosong)
    set_cell_bg(tbl.cell(0, 0), "1a4d8f")

    # Header "Prediksi" span atas
    hdr_pred = tbl.cell(0, 1)
    hdr_pred.merge(tbl.cell(0, nk + 1))
    set_cell_text(hdr_pred, "Prediksi", bold=True, color_rgb=(255, 255, 255))
    set_cell_bg(tbl.cell(0, 1), "1a4d8f")

    # Row header "True \ Pred"
    set_cell_text(tbl.cell(1, 0), "True \\ Pred", bold=True, color_rgb=(255, 255, 255))
    set_cell_bg(tbl.cell(1, 0), "5b87c5")

    # Header kolom kelas
    for j, cls in enumerate(classes):
        cell = tbl.cell(1, j + 1)
        set_cell_text(cell, cls, bold=True, color_rgb=(255, 255, 255))
        set_cell_bg(cell, "5b87c5")

    # Header "Total Aktual"
    set_cell_text(tbl.cell(1, nk + 1), "Total Aktual", bold=True, color_rgb=(255, 255, 255))
    set_cell_bg(tbl.cell(1, nk + 1), "5b87c5")

    # Isi confusion matrix
    for i, cls in enumerate(classes):
        row_idx = i + 2
        set_cell_text(tbl.cell(row_idx, 0), cls, bold=True)
        set_cell_bg(tbl.cell(row_idx, 0), "dde9f7")

        for j in range(nk):
            val = cm[i][j]
            bg = "b8d4ef" if i == j else "f5f8ff"
            set_cell_text(tbl.cell(row_idx, j + 1), str(val), bold=(i == j))
            set_cell_bg(tbl.cell(row_idx, j + 1), bg)

        row_sum = int(cm[i, :].sum())
        set_cell_text(tbl.cell(row_idx, nk + 1), str(row_sum), bold=True)
        set_cell_bg(tbl.cell(row_idx, nk + 1), "e8f0fb")

    doc.add_paragraph()

    # ── SUBTITLE F1 ─────────────────────────────────────────
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub2.add_run("Tabel 2. F1-Score per Kelas Kendaraan")
    r2.bold = True
    r2.font.name = "Times New Roman"

    col_hdrs = ["Kelas Kendaraan", "Precision", "Recall", "F1-Score", "Support"]
    tbl2 = doc.add_table(rows=len(classes) + 3, cols=5)
    tbl2.style = "Table Grid"

    for j, h in enumerate(col_hdrs):
        cell = tbl2.cell(0, j)
        set_cell_text(cell, h, bold=True, color_rgb=(255, 255, 255))
        set_cell_bg(cell, "1a4d8f")

    for i, cls in enumerate(classes):
        m = metrics[cls]
        vals = [cls, f"{m['Precision']:.4f}", f"{m['Recall']:.4f}",
                f"{m['F1']:.4f}", str(support[cls])]
        bg = "dde9f7" if i % 2 == 0 else "f5f8ff"
        for j, v in enumerate(vals):
            cell = tbl2.cell(i + 1, j)
            set_cell_text(cell, v, bold=(j == 3))
            set_cell_bg(cell, bg)

    # Macro Avg
    macro_vals = ["Macro Avg", "-", "-", f"{macro_f1:.4f}", str(total_support)]
    for j, v in enumerate(macro_vals):
        cell = tbl2.cell(len(classes) + 1, j)
        set_cell_text(cell, v, bold=(j == 3))
        set_cell_bg(cell, "c9dcf0")

    # Weighted Avg
    wt_vals = ["Weighted Avg", "-", "-", f"{weighted_f1:.4f}", str(total_support)]
    for j, v in enumerate(wt_vals):
        cell = tbl2.cell(len(classes) + 2, j)
        set_cell_text(cell, v, bold=(j == 3))
        set_cell_bg(cell, "b0cceb")

    doc.add_paragraph()

    # ── RINGKASAN ──────────────────────────────────────────
    ket_hdr = doc.add_paragraph()
    r_ket = ket_hdr.add_run("Ringkasan Hasil Pengujian:")
    r_ket.bold = True
    r_ket.font.name = "Times New Roman"

    ringkasan_items = [
        f"Overall Accuracy       : {overall_acc*100:.2f}%",
        f"Macro F1-Score         : {macro_f1:.4f}",
        f"Weighted F1-Score      : {weighted_f1:.4f}",
        f"Jumlah data uji        : {N} sampel",
        f"Terklasifikasi benar   : {int(cm.diagonal().sum())} sampel",
        f"Misclassification      : {N - int(cm.diagonal().sum())} sampel "
        f"(Pajero SUV→MPV, Grandlivina MPV→City Car)",
    ]
    for item in ringkasan_items:
        p = doc.add_paragraph(style="List Bullet")
        p.text = item
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

    doc.save(filename)
    print(f"[OK] Word document disimpan: {filename}")


# ============================================================
# MAIN
# ============================================================
if __name__ == "__main__":
    plot_confusion_matrix(cm, classes)
    plot_f1_table(metrics, classes, support, macro_f1, weighted_f1, overall_acc)
    export_word(cm, classes, metrics, support, macro_f1, weighted_f1, overall_acc)
    print("\n=== SELESAI ===")
    print("File yang dihasilkan:")
    print("  1. Confusion_Matrix_F1.png")
    print("  2. Tabel_F1_Score.png")
    print("  3. Confusion_Matrix_F1.docx")
