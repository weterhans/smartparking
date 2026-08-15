import math
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, color_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

# Data evaluasi untuk K=3
# Berdasarkan hasil perhitungan sebelumnya:
# Total data = 24. Benar = 22. Accuracy Total = 91.67%
# CityCar : TP=5, TN=18, FP=1, FN=0
# Sedan   : TP=4, TN=20, FP=0, FN=0
# MPV     : TP=4, TN=18, FP=1, FN=1
# SUV     : TP=4, TN=19, FP=0, FN=1
# Pickup  : TP=5, TN=19, FP=0, FN=0

data_k3 = [
    {"label": "CityCar", "TP": 5, "TN": 18, "FP": 1, "FN": 0},
    {"label": "Sedan",   "TP": 4, "TN": 20, "FP": 0, "FN": 0},
    {"label": "MPV",     "TP": 4, "TN": 18, "FP": 1, "FN": 1},
    {"label": "SUV",     "TP": 4, "TN": 19, "FP": 0, "FN": 1},
    {"label": "Pickup",  "TP": 5, "TN": 19, "FP": 0, "FN": 0}
]

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Tabel Hasil Pengujian Model dengan Nilai K = 3")
run.bold = True

table = doc.add_table(rows=6, cols=5)
table.style = 'Table Grid'

headers = ["Label Kendaraan", "Precision (%)", "Recall (%)", "Accuracy (%)", "Accuracy Total (%)"]
for i, h in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = h
    set_cell_background(cell, "A6A6A6")  # Warna abu-abu seperti screenshot
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Fill data
for i, d in enumerate(data_k3):
    row_idx = i + 1
    
    precision = (d["TP"] / (d["TP"] + d["FP"])) * 100 if (d["TP"] + d["FP"]) > 0 else 0
    recall = (d["TP"] / (d["TP"] + d["FN"])) * 100 if (d["TP"] + d["FN"]) > 0 else 0
    accuracy_kelas = ((d["TP"] + d["TN"]) / 24) * 100
    
    # Format to string without decimal if it's a whole number, else 1 decimal
    str_prec = f"{int(precision)}%" if precision.is_integer() else f"{precision:.1f}%"
    str_rec = f"{int(recall)}%" if recall.is_integer() else f"{recall:.1f}%"
    str_acc = f"{int(accuracy_kelas)}%" if accuracy_kelas.is_integer() else f"{accuracy_kelas:.1f}%"
    
    cells = table.rows[row_idx].cells
    cells[0].text = d["label"]
    cells[1].text = str_prec
    cells[2].text = str_rec
    cells[3].text = str_acc
    
    # Center align text
    for j in range(4):
        cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Merge the last column for total accuracy
a = table.cell(1, 4)
b = table.cell(5, 4)
a.merge(b)

merged_cell = table.cell(1, 4)
merged_cell.text = "91.67%"
merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

out_file = "Tabel_Evaluasi_K3.docx"
doc.save(out_file)
print(f"File disimpan sebagai {out_file}")
