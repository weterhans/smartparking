import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Paragraf pembuka
p1 = doc.add_paragraph('Proses penentuan kelas kendaraan baru pada tahap klasifikasi menggunakan algoritma ')
p1.add_run('Feature-Weighted K-Nearest Neighbor').italic = True
p1.add_run(' (FWk-NN) dilakukan melalui langkah-langkah sebagai berikut:')
p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# List items
items = [
    'Sistem memulai proses klasifikasi dengan membaca data mentah dari sensor TF-LUNA LiDAR. Data profil bentuk kendaraan ini kemudian diekstraksi menjadi 8 fitur fisis utama, yaitu: Panjang (L), Tinggi (H), StdDev, PosMax, Slope, Compactness, RearComp, dan FlatRoof.',
    'Kedelapan fitur mentah dari kendaraan yang baru dideteksi (data uji) kemudian dinormalisasi menggunakan persamaan Min-Max Scaling. Proses ini mereduksi seluruh nilai fitur menjadi skala seragam antara 0 hingga 1 menggunakan nilai minimum dan maksimum dari dataset latih sebagai acuan.',
    'Nilai fitur data uji yang telah dinormalisasi kemudian dikalikan dengan parameter bobot optimal yang telah ditetapkan (merujuk pada Tabel X).',
    'Sistem menghitung tingkat kedekatan antara data uji dengan 75 data referensi (data latih) menggunakan persamaan Weighted Euclidean Distance. Rumus perhitungan jarak tersebut adalah sebagai berikut:'
]

for i, text in enumerate(items, 1):
    p = doc.add_paragraph(f"{i}. {text}")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Formula 1
p_f1 = doc.add_paragraph('D(x,y) = √ ∑ [ w_i × (x_i - y_i)² ]')
p_f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_f1.runs[0].bold = True

p_desc1 = doc.add_paragraph('Di mana D(x,y) adalah jarak antara data uji x dan data latih y, w_i adalah bobot fitur ke-i, serta x_i dan y_i adalah nilai fitur ke-i yang telah dinormalisasi.')
p_desc1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_desc1.paragraph_format.left_indent = Pt(18)

# Item 5
p5 = doc.add_paragraph('5. Dari hasil kalkulasi jarak tersebut, sistem melakukan penyortiran secara terurut (ascending sort) untuk mengekstraksi sejumlah K data latih yang memiliki jarak minimum (tingkat kemiripan tertinggi) terhadap data uji. Pada implementasi akhir sistem ini, parameter K ditetapkan bernilai 3 (K=3).')
p5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Item 6
p6 = doc.add_paragraph('6. Setelah 3 tetangga terdekat diidentifikasi, klasifikasi akhir ditentukan menggunakan metode ')
p6.add_run('Majority Voting').italic = True
p6.add_run(' (pemungutan suara mayoritas) terhadap label kelas dari ketiga data latih tersebut, dengan menggunakan rumus berikut:')
p6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Formula 2
p_f2 = doc.add_paragraph("y' = argmax_v ∑ I(y_i = v)")
p_f2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_f2.runs[0].bold = True

p_desc2 = doc.add_paragraph("Di mana y' adalah kelas prediksi akhir, v adalah kandidat label kelas, I(·) adalah fungsi indikator yang bernilai 1 jika kondisi benar (tetangga ke-i memiliki kelas v) dan 0 jika salah.")
p_desc2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_desc2.paragraph_format.left_indent = Pt(18)

# Paragraf penutup
p_end = doc.add_paragraph('Sebagai ilustrasi, apabila dari ketiga tetangga terdekat (K=3) didapatkan 2 (dua) data berlabel kelas City Car dan 1 (satu) data berlabel kelas Sedan, maka berdasarkan frekuensi kemunculan tertinggi, sistem secara konklusif memprediksi dan mengklasifikasikan kendaraan uji tersebut sebagai City Car.')
p_end.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.save('Penjelasan_Klasifikasi_FWkNN.docx')
print("File Word berhasil dibuat!")
